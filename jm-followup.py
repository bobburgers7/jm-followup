import pandas as pd
from docx import Document
import datetime
import os
import math

# remember to relabel the CITY/STATE columns in the excel sheet - they are backwards
# ADDRESS LINE 2 has a SPACE in front of it

currYear = datetime.date.today().strftime("%Y")
currMonth = datetime.date.today().strftime("%m")

def is_odd(a):
    return bool(a & 1)

# https://github.com/python-openxml/python-docx/issues/33#issuecomment-77661907
# erases the extra paragraph in the table making it take more space
def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

#spit out a word doc
def create_word_doc(specialties, dataFrame, headerOfFile):
    # create word document
    document = Document()
    document.add_heading(headerOfFile, 0)

    for specialty in specialties:
        physicians = dataFrame[dataFrame.SPECIALTY1 == specialty]
        # get rid of empty fields and fill w/ empty string
        physicians = physicians.fillna('')

        document.add_heading(specialty, 1)

        # find out how many rows I need as an integer; add an extra row if # is odd
        rowCount = physicians.shape[0]  # gives number of rows in this dataframe
        rowNum: int = int(math.ceil(rowCount / 2))
        # add table (rows, columns)
        table = document.add_table(rowNum, 2)
        currColumn = 0
        currRow = 0

        for index, row in physicians.iterrows():
            physicianName = row['LAST NAME'] + ", " + row['FIRST NAME'] + " " + row['TITLE']
            officeAddress = row['ADDRESS LINE1'] + " " + row[' ADDRESS LINE 2'] + "\n" + row['CITY'] + ', ' + row[
                'STATE'] + ' ' + str(row['ZIP'])
            officePhone = row['OFFICE PHONE']
            officeFax = row['OFFICE FAX']

            cell = table.cell(currRow, currColumn)
            delete_paragraph(cell.paragraphs[-1])

            p = cell.add_paragraph()
            p.add_run(physicianName + "\n").bold = True
            p.add_run(officeAddress + "\n")
            p.add_run("Office Phone: " + officePhone + "\n")
            p.add_run("Office Fax: " + officeFax)

            if is_odd(currColumn):
                currColumn = 0
                currRow = currRow + 1
            else:
                currColumn = 1

    document.add_paragraph('Get the most up to date information and schedule an appointment online: https://www.johnmuirhealth.com/fad/')

    fileName = headerOfFile + '.docx'
    filePath = os.path.join('worddoc', fileName)
    document.save(filePath)

# don't specify sheet name so can just get the 1st
excelSheet = pd.read_excel('jmdirectory2018.xls', index_col=None, na_values=['NA'])

# drop e-mail column; Note: axis=1 denotes that we are referring to a column, not a row
excelSheet = excelSheet.drop('EMAIL ADDRESS', axis=1)

# remove rows containing PAs, certified nursing midwife, NP
titlesToRemove = ('PA', 'PA-C', 'CNM', 'NP', 'RN', 'RNFA')
for title in titlesToRemove:
    excelSheet = excelSheet[excelSheet.TITLE != title]

#combine specialties that should show up together
thoracicSpecialities = ('Cardiac Surgery', 'Cardiothoracic Surgery', 'Thoracic Surgery')
cardiologySubSpecialities = ('Cardiology', 'Cardiac Electrophysiology', 'Interventional Cardiology')
ophthoSubSpecialties = ('Ophthalmology', 'Oculoplastic Surgery', 'Retinal Ophthalmology')
surgerySubSpecialties = ('General Surgery', 'Colon and Rectal Surgery')
obSubSpecialties = ('Obstetrics and Gynecology', 'Perinatology', 'Gynecologic Oncology', 'Gynecology', 'Obstetrics', 'Reproductive Endocrinology and Infertility')

# carve out pediatric subspecialties note the space after Pediatric which will collect Peds cards, etc.
pediatricSubspecialities = excelSheet[excelSheet.SPECIALTY1.str.contains("Pediatric ") == 1][
    'SPECIALTY1'].drop_duplicates()

create_word_doc(thoracicSpecialities, excelSheet, "John Muir - Thoracic & Cardiothoracic Surgery")
create_word_doc(cardiologySubSpecialities, excelSheet, 'John Muir - Cardiology')
create_word_doc(ophthoSubSpecialties, excelSheet, "John Muir - Ophthalmology")
create_word_doc(surgerySubSpecialties, excelSheet, 'John Muir - General Surgery and Colorectal Surgery')
create_word_doc(obSubSpecialties, excelSheet, 'John Muir - OB-GYN and Gyn-Onc')
create_word_doc(pediatricSubspecialities, excelSheet, 'John Muir - Pediatric Specialties')

specialtiesToRemove = ('Addiction Specialist',
                       'Anesthesiology',
                       'Cardiac Anesthesiology',
                       'Dentistry',
                       'Diagnostic Radiology',
                       'Emergency Medicine',
                       'Gastroenterology (Hospital-Based Only)',
                       'General Surgery-Surgical Assist',
                       'Hospitalist',
                       'Hyperbaric Medicine',
                       'Neonatology',
                       'Nurse Practitioner - Breast Health',
                       'Nurse Practitioner - Palliative Care',
                       'Palliative Care',
                       'Pathology',
                       'Pediatric Hospitalist',
                       'Pediatric Radiology',
                       'Perioperative Medicine',
                       'Physician Assistant - Orthopedic',
                       'Registered Nurse First Assist (RNFA)',
                       'Spine Specialist',
                       'Surgical Assistant',
                       'Teleradiology',
                       'Urgent Care Provider')

#remove all the specialties I've already made lists of
specialtiesToRemove = specialtiesToRemove + thoracicSpecialities + cardiologySubSpecialities + ophthoSubSpecialties + surgerySubSpecialties + obSubSpecialties

for specialty in specialtiesToRemove:
    excelSheet = excelSheet[excelSheet.SPECIALTY1 != specialty]

# cuts out all the pediatric subspecialists + duplicates
specialties = excelSheet[excelSheet.SPECIALTY1.str.contains("Pediatric ") == 0]['SPECIALTY1'].drop_duplicates()

for specialty in specialties:
    # list all physician of X specialty
    physicians = excelSheet[excelSheet.SPECIALTY1 == specialty]

    # get rid of empty fields
    physicians = physicians.fillna('')

    # create new word document
    document = Document()
    document.add_heading('John Muir - ' + specialty, 0)

    # find out how many rows I need as an integer; add an extra row if # is odd
    rowCount = physicians.shape[0]  # gives number of rows in this dataframe
    rowNum = int(math.ceil(rowCount / 2))
    # add table (rows, columns)
    table = document.add_table(rowNum, 2)
    currColumn = 0
    currRow = 0

    for index, row in physicians.iterrows():
        physicianName = row['LAST NAME'] + ", " + row['FIRST NAME'] + " " + row['TITLE']
        officeAddress = row['ADDRESS LINE1'] + " " + row[' ADDRESS LINE 2'] + "\n" + row['CITY'] + ', ' + row[
            'STATE'] + ' ' + str(row['ZIP'])
        officePhone = row['OFFICE PHONE']
        officeFax = row['OFFICE FAX']

        cell = table.cell(currRow, currColumn)
        delete_paragraph(cell.paragraphs[-1])

        p = cell.add_paragraph()
        p.add_run(physicianName + "\n").bold = True
        p.add_run(officeAddress + "\n")
        p.add_run("Office Phone: " + officePhone + "\n")
        p.add_run("Office Fax: " + officeFax)

        if is_odd(currColumn):
            currColumn = 0
            currRow = currRow + 1
        else:
            currColumn = 1

    document.add_paragraph('Get the most up to date information and schedule an appointment online: https://www.johnmuirhealth.com/fad/')

    # deal w/ heme/onc creating a new directory
    specialty = specialty.replace("/", "-")
    fileName = specialty + '.docx'
    filePath = os.path.join('worddoc', fileName)
    document.save(filePath)

# https://stackoverflow.com/questions/16476924/how-to-iterate-over-rows-in-a-dataframe-in-pandas

# deal with blank fields by making them empty strings
# https://stackoverflow.com/questions/29782898/combine-pandas-data-frame-column-values-into-new-column